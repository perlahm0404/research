#!/usr/bin/env python3
"""
Convert .docx files to Markdown format
Requires: python-docx library
Install: pip install python-docx
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

def extract_text_from_paragraph(para):
    """Extract text from a paragraph, preserving formatting hints"""
    text = ""
    for run in para.runs:
        if run.bold:
            text += f"**{run.text}**"
        elif run.italic:
            text += f"*{run.text}*"
        else:
            text += run.text
    return text

def convert_docx_to_md(docx_path, output_path=None):
    """
    Convert a DOCX file to Markdown

    Args:
        docx_path: Path to the .docx file
        output_path: Path for the output .md file (optional)

    Returns:
        Path to the created markdown file
    """

    docx_path = Path(docx_path)

    if not docx_path.exists():
        print(f"Error: File not found: {docx_path}")
        return None

    if not docx_path.suffix.lower() == '.docx':
        print(f"Error: File is not a .docx file: {docx_path}")
        return None

    # Set output path
    if output_path is None:
        output_path = docx_path.with_suffix('.md')
    else:
        output_path = Path(output_path)

    print(f"Converting: {docx_path.name}")
    print(f"Output: {output_path.name}")

    try:
        # Load the document
        doc = Document(docx_path)

        # Build markdown content
        md_content = []

        # Add title from filename if no heading found
        md_content.append(f"# {docx_path.stem.replace('-', ' ')}\n")
        md_content.append(f"*Converted from: {docx_path.name}*\n\n")

        # Process paragraphs
        for para in doc.paragraphs:
            text = extract_text_from_paragraph(para)

            # Skip empty paragraphs
            if text.strip():
                # Check for heading styles
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                    md_content.append(f"{'#' * level} {text}\n\n")
                else:
                    md_content.append(f"{text}\n\n")

        # Process tables
        for table in doc.tables:
            md_content.append("\n")
            # Build markdown table
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                md_content.append("| " + " | ".join(cells) + " |\n")

                # Add header separator after first row
                if i == 0:
                    md_content.append("|" + "|".join(["---"] * len(cells)) + "|\n")
            md_content.append("\n")

        # Write to markdown file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(''.join(md_content))

        print(f"✓ Successfully converted to: {output_path}")
        return output_path

    except Exception as e:
        print(f"✗ Error converting {docx_path}: {e}")
        import traceback
        traceback.print_exc()
        return None

def convert_all_docx_in_directory(directory, pattern="*.docx"):
    """
    Convert all .docx files in a directory to markdown

    Args:
        directory: Path to directory containing .docx files
        pattern: File pattern to match (default: "*.docx")

    Returns:
        List of created markdown files
    """

    directory = Path(directory)

    if not directory.exists():
        print(f"Error: Directory not found: {directory}")
        return []

    docx_files = list(directory.glob(pattern))

    if not docx_files:
        print(f"No .docx files found in {directory}")
        return []

    print(f"Found {len(docx_files)} .docx file(s) to convert\n")

    created_files = []
    for docx_file in docx_files:
        md_file = convert_docx_to_md(docx_file)
        if md_file:
            created_files.append(md_file)
        print()

    print(f"\nConversion complete! Created {len(created_files)} markdown file(s)")
    return created_files

if __name__ == "__main__":
    if len(sys.argv) > 1:
        # Convert specific file or directory
        path = sys.argv[1]
        path = Path(path)

        if path.is_file() and path.suffix.lower() == '.docx':
            convert_docx_to_md(path)
        elif path.is_dir():
            convert_all_docx_in_directory(path)
        else:
            print(f"Error: Invalid path or not a .docx file: {path}")
    else:
        # Convert all .docx files in current directory
        print("Usage: python convert_docx_to_md.py [file.docx or directory]")
        print("\nRunning conversion on current directory...\n")
        convert_all_docx_in_directory(Path.cwd())